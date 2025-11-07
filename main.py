import os
import re
from flask import Flask, request, redirect, url_for, render_template_string, session, flash, send_from_directory
from werkzeug.security import check_password_hash, generate_password_hash
from openpyxl import Workbook, load_workbook
from docx import Document
from datetime import datetime

# ---------------- CONFIG ----------------
SITE_NAME = "TamEcoVita Suites"
APP_SECRET = os.environ.get("APP_SECRET", "change_this_secret_for_prod")
ADMIN_PASSWORD_HASH = os.environ.get(
    "ADMIN_PASSWORD_HASH",
    generate_password_hash(os.environ.get("ADMIN_PASSWORD", "tamecovita1"))
)

DATA_FOLDER = os.path.join(os.getcwd(), "instance")
os.makedirs(DATA_FOLDER, exist_ok=True)

FILENAME = os.path.join(DATA_FOLDER, "TamEcoVita_host_file.xlsx")
CUSTOMER_FOLDER = os.path.join(DATA_FOLDER, "requests")
os.makedirs(CUSTOMER_FOLDER, exist_ok=True)

app = Flask(__name__)
app.secret_key = APP_SECRET

@app.route('/logo.png')
def serve_logo():
    return send_from_directory(os.getcwd(), 'logo.png')

# ---------------- HELPERS ----------------
def ensure_excel():
    if not os.path.exists(FILENAME):
        wb = Workbook()
        ws = wb.active
        ws.title = "Reservations"
        ws.append([
            "Receipt No", "Guest Name", "Contact/Email", "Apartment Type",
            "Place / Location", "Check-In Date", "Check-Out Date",
            "Number of Nights", "Rate per Night", "VAT", "Total Amount",
            "Amount Paid", "Payment Method", "Payment Date", "Balance"
        ])
        wb.save(FILENAME)
ensure_excel()

def sanitize_filename(name):
    return re.sub(r'[<>:"/\\|?*]', '_', name)

def next_id_code(ws):
    count = ws.max_row - 1
    return f"TEC-{count+1:04d}"

def auto_adjust_columns(ws):
    for column in ws.columns:
        if not column[0].column_letter:
            continue
        max_len = max((len(str(cell.value)) if cell.value else 0) for cell in column)
        ws.column_dimensions[column[0].column_letter].width = max_len + 5


# ---------------- TEMPLATES ----------------
login_template = """<!doctype html>
<title>Admin Login - {{ SITE_NAME }}</title>
<style>
body {
    font-family: Arial, sans-serif;
    background-color: #f5f7fa;
    display: flex;
    justify-content: center;
    align-items: center;
    height: 100vh;
    margin: 0;
}
.login-container {
    background: white;
    padding: 40px;
    border-radius: 10px;
    box-shadow: 0 4px 20px rgba(0,0,0,0.15);
    width: 350px;
    text-align: center;
}
.logo {
    width: 120px;
    margin-bottom: 20px;
}
h2 {
    color: #1b2a49;
    margin-bottom: 20px;
}
input[type="password"] {
    width: 100%;
    padding: 10px;
    margin-bottom: 15px;
    border-radius: 5px;
    border: 1px solid #ccc;
}
input[type="submit"] {
    width: 100%;
    background-color: #1b2a49;
    color: white;
    padding: 10px;
    border: none;
    border-radius: 6px;
    font-weight: bold;
    cursor: pointer;
    transition: 0.2s;
}
input[type="submit"]:hover {
    background-color: #304269;
}
.message {
    color: red;
    margin-top: 10px;
}
</style>

<div class="login-container">
    <img src="/logo.png" class="logo" alt="Logo">
    <h2>Admin Login</h2>
    <form method="post">
        <input type="password" name="password" placeholder="Enter Password" required>
        <input type="submit" value="Login">
    </form>
    {% for message in get_flashed_messages() %}
    <p class="message">{{ message }}</p>
    {% endfor %}
</div>
"""

index_template = """<!doctype html>
<title>{{ SITE_NAME }} - Admin Dashboard</title>
<style>
body {
    font-family: Arial, sans-serif;
    background: #f5f7fa;
    margin: 0;
    padding: 0;
}
header {
    background: #1b2a49;
    color: white;
    padding: 15px 0;
    text-align: center;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}
header img {
    width: 100px;
    display: block;
    margin: 0 auto 10px auto;
}
header h2 {
    margin: 0;
    font-size: 24px;
}
.container {
    width: 90%;
    margin: 30px auto;
    background: white;
    padding: 20px 30px;
    border-radius: 10px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.1);
}
h3 {
    color: #1b2a49;
    margin-bottom: 15px;
}
table {
    width: 100%;
    border-collapse: collapse;
    margin-top: 10px;
}
th, td {
    padding: 10px;
    border: 1px solid #ccc;
    text-align: left;
}
th {
    background-color: #1b2a49;
    color: white;
}
a {
    color: #1b2a49;
    text-decoration: none;
}
a:hover {
    text-decoration: underline;
}
.actions a {
    margin-right: 10px;
}
.top-links {
    text-align: center;
    margin-bottom: 20px;
}
.top-links a {
    background-color: #1b2a49;
    color: white;
    padding: 10px 15px;
    border-radius: 5px;
    margin: 0 5px;
    text-decoration: none;
    transition: background-color 0.2s;
}
.top-links a:hover {
    background-color: #304269;
}
</style>

<header>
    <img src="/logo.png" alt="Logo">
    <h2>{{ SITE_NAME }} - Admin Dashboard</h2>
</header>

<div class="container">
<div class="top-links">
<a href="{{ url_for('logout') }}">Logout</a>
<a href="{{ url_for('download') }}">Download Excel</a>
<a href="{{ url_for('search') }}">Search</a>
</div>
<h3>Reservations</h3>
<table>
<tr><th>Receipt</th><th>Guest</th><th>Check-In</th><th>Check-Out</th><th>Actions</th></tr>
{% for row in rows %}
<tr>
<td>{{ row['Receipt No'] }}</td>
<td>{{ row['Guest Name'] }}</td>
<td>{{ row['Check-In Date'] }}</td>
<td>{{ row['Check-Out Date'] }}</td>
<td class="actions">
<a href="{{ url_for('edit', receipt_no=row['Receipt No']) }}">Edit</a> |
<a href="{{ url_for('download_word_receipt', receipt_no=row['Receipt No']) }}">Receipt</a>|
<a href="{{ url_for('download_invoice', receipt_no=row['Receipt No']) }}">Invoice</a>
</td>
</tr>
{% endfor %}
</table>
</div>
"""

edit_template = """<!doctype html><title>Edit Reservation</title>
<h2>Edit Reservation {{ record['Receipt No'] }}</h2>
<form method="post" action="{{ url_for('update', receipt_no=record['Receipt No']) }}">
{% for key, value in record.items() %}
<label>{{ key }}:</label> <input name="{{ key }}" value="{{ value }}"><br>
{% endfor %}
<input type="submit" value="Update">
</form>
<a href="{{ url_for('index') }}">Back</a>
"""

search_template = """<!doctype html><title>Search Reservations</title>
<h2>Search</h2>
<form method="get">
<input type="text" name="q" placeholder="Search by guest name">
<input type="submit" value="Search">
</form>
<table border="1">
<tr><th>Receipt</th><th>Guest</th><th>Check-In</th><th>Check-Out</th></tr>
{% for row in rows %}
<tr>
<td>{{ row['Receipt No'] }}</td>
<td>{{ row['Guest Name'] }}</td>
<td>{{ row['Check-In Date'] }}</td>
<td>{{ row['Check-Out Date'] }}</td>
</tr>
{% endfor %}
</table>
<a href="{{ url_for('index') }}">Back</a>
"""

# ---------------- UTILITY ----------------
def fill_word_template(template_path, output_path, mapping):
    """Replaces placeholders in a Word file based on a dict mapping."""
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if f"{{{key}}}" in p.text:
                p.text = p.text.replace(f"{{{key}}}", str(val))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in mapping.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", str(val))
    doc.save(output_path)
    return output_path


def generate_invoice(receipt_no, record):
    template_path = os.path.join(DATA_FOLDER, "invoice_template.docx")
    guest_clean = sanitize_filename(record.get("Guest Name", "Guest"))
    output_path = os.path.join(DATA_FOLDER, f"Invoice_{guest_clean}.docx")

    mapping = {
        "DATE": datetime.today().strftime("%Y-%m-%d"),
        "INVOICE_NO": receipt_no,
        "GUEST_NAME": record.get("Guest Name", ""),
        "APARTMENT": record.get("Apartment Type", ""),
        "UNIT_PRICE": record.get("Rate per Night", ""),
        "DAYS": record.get("Number of Nights", ""),
        "SUBTOTAL": f"{float(record.get('Rate per Night', 0)) * float(record.get('Number of Nights', 0)):.2f}",
        "VAT": record.get("VAT", ""),
        "TOTAL_DUE": record.get("Total Amount", "")
    }

    return fill_word_template(template_path, output_path, mapping)


def generate_word_receipt(receipt_no, record):
    template_path = os.path.join(DATA_FOLDER, "receipt_template.docx")
    guest_clean = sanitize_filename(record.get("Guest Name", "Guest"))
    output_path = os.path.join(DATA_FOLDER, f"Receipt_{guest_clean}.docx")

    mapping = {
        "GUEST_NAME": record.get("Guest Name", ""),
        "CONTACT": record.get("Contact/Email", ""),
        "DATE": datetime.today().strftime("%Y-%m-%d"),
        "RECEIPT_NO": receipt_no,
        "BOOKING_REF": receipt_no,
        "APARTMENT_TYPE": record.get("Apartment Type", ""),
        "LOCATION": record.get("Place / Location", ""),
        "CHECKIN": record.get("Check-In Date", ""),
        "CHECKOUT": record.get("Check-Out Date", ""),
        "NIGHTS": record.get("Number of Nights", ""),
        "RATE": record.get("Rate per Night", ""),
        "VAT": record.get("VAT", ""),
        "TOTAL": record.get("Total Amount", ""),
        "PAID": record.get("Amount Paid", ""),
        "METHOD": record.get("Payment Method", ""),
        "PAYMENT_DATE": record.get("Payment Date", ""),
        "BALANCE": record.get("Balance", "")
    }

    return fill_word_template(template_path, output_path, mapping)

# ---------------- AUTH ----------------
def logged_in():
    return session.get("logged_in") == True

@app.route("/login", methods=["GET","POST"])
def login():
    if request.method == "POST":
        if check_password_hash(ADMIN_PASSWORD_HASH, request.form.get("password","")):
            session["logged_in"] = True
            return redirect(url_for("index"))
        flash("Incorrect password.")
    return render_template_string(login_template, SITE_NAME=SITE_NAME)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------------- HOME ----------------
@app.route("/home")
@app.route("/home")
def home():
    home_template = f"""
    <!doctype html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>{SITE_NAME} - Home</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f5f7fa;
                text-align: center;
                margin: 0;
                padding: 0;
                display: flex;
                flex-direction: column;
                min-height: 100vh;
            }}
            main {{
                flex: 1;
            }}
            .logo {{
                margin-top: 50px;
                width: 180px;
            }}
            h1 {{
                color: #1b2a49;
                margin-top: 20px;
            }}
            .btn {{
                display: inline-block;
                margin: 20px;
                padding: 15px 30px;
                background-color: #1b2a49;
                color: white;
                border-radius: 8px;
                text-decoration: none;
                font-weight: bold;
            }}
            .btn:hover {{
                background-color: #304269;
            }}
            footer {{
                background-color: #fff;
                color: #333;
                padding: 25px;
                text-align: center;
                box-shadow: 0 -2px 8px rgba(0,0,0,0.1);
                font-size: 15px;
                line-height: 1.6;
            }}
            footer a {{
                color: #1b2a49;
                text-decoration: none;
            }}
            footer a:hover {{
                text-decoration: underline;
            }}
        </style>
    </head>
    <body>
        <main>
            <img src="/logo.png" class="logo">
            <h1>Welcome to {SITE_NAME}</h1>
            <a href="/customer" class="btn">Customer</a>
            <a href="/login" class="btn">Admin</a>
        </main>

        <footer>
            <strong>TAM Ecovista Properties</strong><br>
            Tel: +2348117759059, +2348135567475<br>
            RC: 7405900<br>
            Address: 27B First Avenue, Gwarinpa, Abuja<br>
            Website: <a href="https://tamecovista.com" target="_blank">tamecovista.com</a><br>
            Email: <a href="mailto:info@tamecovista.com">info@tamecovista.com</a>
        </footer>
    </body>
    </html>
    """
    return home_template




@app.route("/")
def root_redirect():
    return redirect(url_for("home"))

# ---------------- CUSTOMER ----------------
customer_template = """<!doctype html>
<title>Reservation Request</title>
<style>
body {
    font-family: Arial, sans-serif;
    background: #f5f7fa;
    margin: 0;
    padding: 0;
}
.container {
    max-width: 800px;
    margin: 50px auto;
    padding: 30px;
    background: white;
    border-radius: 10px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.1);
}
.logo {
    display: block;
    margin: 0 auto 20px auto;
    width: 150px;
}
h2 {
    text-align: center;
    color: #1b2a49;
    margin-bottom: 30px;
}
.form-grid {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 20px;
}
label {
    display: block;
    margin-bottom: 5px;
    font-weight: bold;
    color: #333;
}
input, select {
    width: 100%;
    padding: 8px;
    border-radius: 5px;
    border: 1px solid #ccc;
}
input[type="submit"] {
    grid-column: 1 / -1;
    background: #1b2a49;
    color: white;
    padding: 12px;
    font-size: 16px;
    border: none;
    border-radius: 6px;
    cursor: pointer;
    margin-top: 20px;
    transition: 0.2s;
}
input[type="submit"]:hover {
    background: #304269;
}
.message {
    color: green;
    text-align: center;
    margin-top: 15px;
}
</style>

<div class="container">
<img src="/logo.png" class="logo" alt="Logo">
<h2>Reservation Request Form</h2>

<form method="post" action="{{ url_for('submit') }}">
<div class="form-grid">
    <!-- Left column: Guest Info -->
    <div>
        <label>Guest Name:</label><input name="name" required>
        <label>Contact/Email:</label><input name="email" type="email">
        <label>Number Of Nights:</label><input name="total_days" type="number" step="1">
        <label>Check-In Date:</label><input name="date_coming" type="date" required>
        <label>Check-Out Date:</label><input name="date_going" type="date" required>
        <label>Place / Location:</label>
        <select name="place">
    <option value="">--Select Location--</option>

    <option value="Abaji">Abaji</option>
    <option value="Aco Estate">Aco Estate</option>
    <option value="Airport Road">Airport Road</option>
    <option value="Apo">Apo</option>
    <option value="Apo Dutse">Apo Dutse</option>
    <option value="Area 1">Area 1</option>
    <option value="Area 11">Area 11</option>
    <option value="Area 3">Area 3</option>
    <option value="Asokoro">Asokoro</option>
    <option value="Bude">Bude</option>
    <option value="Bunkoro">Bunkoro</option>
    <option value="Burun">Burun</option>
    <option value="Bwari">Bwari</option>
    <option value="Central Area">Central Area</option>
    <option value="Chafe">Chafe</option>
    <option value="Chika">Chika</option>
    <option value="City Centre">City Centre</option>
    <option value="Dape">Dape</option>
    <option value="Dakibiyu">Dakibiyu</option>
    <option value="Dakwo">Dakwo</option>
    <option value="Dei-Dei">Dei-Dei</option>
    <option value="Duboyi">Duboyi</option>
    <option value="Durumi">Durumi</option>
    <option value="Dutse Alhaji">Dutse Alhaji</option>
    <option value="Dutse Makaranta">Dutse Makaranta</option>
    <option value="Galadimawa">Galadimawa</option>
    <option value="Gaduwa">Gaduwa</option>
    <option value="Garki">Garki</option>
    <option value="Garki II">Garki II</option>
    <option value="Gbazango">Gbazango</option>
    <option value="Gbazango West">Gbazango West</option>
    <option value="Gidari Bahagwo">Gidari Bahagwo</option>
    <option value="Gosa">Gosa</option>
    <option value="Gudu">Gudu</option>
    <option value="Guzape I">Guzape I</option>
    <option value="Guzape II">Guzape II</option>
    <option value="Gwagwa">Gwagwa</option>
    <option value="Gwagwalada">Gwagwalada</option>
    <option value="Gwarinpa">Gwarinpa</option>
    <option value="Gui">Gui</option>
    <option value="Gwari">Gwari</option>
    <option value="Idu">Idu</option>
    <option value="Idogwari">Idogwari</option>
    <option value="Industrial Area">Industrial Area</option>
    <option value="Jabi">Jabi</option>
    <option value="Jahi">Jahi</option>
    <option value="Jaite">Jaite</option>
    <option value="Kaba">Kaba</option>
    <option value="Kado">Kado</option>
    <option value="Kabusa">Kabusa</option>
    <option value="Kafe">Kafe</option>
    <option value="Kagini">Kagini</option>
    <option value="Kamo">Kamo</option>
    <option value="Karu">Karu</option>
    <option value="Karshi">Karshi</option>
    <option value="Karsana">Karsana</option>
    <option value="Katampe">Katampe</option>
    <option value="Kaura">Kaura</option>
    <option value="Ketti">Ketti</option>
    <option value="Kpoto">Kpoto</option>
    <option value="Kpeyegi">Kpeyegi</option>
    <option value="Kubwa">Kubwa</option>
    <option value="Kuje">Kuje</option>
    <option value="Kuje Hills">Kuje Hills</option>
    <option value="Kukwaba">Kukwaba</option>
    <option value="Kurudu">Kurudu</option>
    <option value="Kurudu Hill">Kurudu Hill</option>
    <option value="Kwali">Kwali</option>
    <option value="Kyami">Kyami</option>
    <option value="Lifecamp">Lifecamp</option>
    <option value="Lokogoma">Lokogoma</option>
    <option value="Lugbe">Lugbe</option>
    <option value="Mabushi">Mabushi</option>
    <option value="Maitama">Maitama</option>
    <option value="Mamusa">Mamusa</option>
    <option value="Mbora">Mbora</option>
    <option value="Mpape">Mpape</option>
    <option value="National stadium">National stadium</option>
    <option value="Nbora">Nbora</option>
    <option value="Nyanya">Nyanya</option>
    <option value="Okanje">Okanje</option>
    <option value="Orozo">Orozo</option>
    <option value="Parfun">Parfun</option>
    <option value="Pegi">Pegi</option>
    <option value="Pyakasa">Pyakasa</option>
    <option value="Sabon Gari">Sabon Gari</option>
    <option value="Sabon Lugbe">Sabon Lugbe</option>
    <option value="Sabo Gida">Sabo Gida</option>
    <option value="Saraji">Saraji</option>
    <option value="Sauka">Sauka</option>
    <option value="Sheretti">Sheretti</option>
    <option value="Suleja">Suleja</option>
    <option value="Tasha">Tasha</option>
    <option value="Tungan Maje">Tungan Maje</option>
    <option value="Utako">Utako</option>
    <option value="Waru-Pozema">Waru-Pozema</option>
    <option value="Wumba">Wumba</option>
    <option value="Wupa">Wupa</option>
    <option value="Wuse">Wuse</option>
    <option value="Wuse II">Wuse II</option>
    <option value="Wuye">Wuye</option>
    <option value="Yimi">Yimi</option>
    <option value="Zuba">Zuba</option>
</select>


    </div>

    <!-- Right column: Apartment & Payment Info -->
    <div>
        <label>Apartment Type:</label><input name="apartment" type="number" step="1">
        <label>Rate Per Night:</label><input name="reservation_fee">
        <label>Payment Method:</label>
        <select name="payment">
            <option value="">--Select Payment Method--</option>
            <option value="Cash">Cash</option>
            <option value="Credit Card">Credit Card</option>
            <option value="Debit Card">Debit Card</option>
            <option value="Bank Transfer">Bank Transfer</option>
            <option value="PayPal">PayPal</option>
            <option value="Google Pay">Google Pay</option>
        </select>
        <label>Payment Date:</label><input name="date" type="date">
        <label>Amount Paid:</label><input name="amount_paid">
    </div>
</div>
<input type="submit" value="Submit Reservation">
</form>

{% for message in get_flashed_messages() %}<p class="message">{{ message }}</p>{% endfor %}
</div>
"""

@app.route("/customer")
def customer_form():
    return render_template_string(customer_template)

@app.route("/submit", methods=["POST"])
def submit():
    def parse_float(value):
        try:
            return float(value)
        except (ValueError, TypeError):
            return 0.0

    name = request.form.get("name", "").strip()
    email = request.form.get("email", "").strip()
    date_coming = request.form.get("date_coming", "").strip()
    date_going = request.form.get("date_going", "").strip()
    place = request.form.get("place", "").strip()
    apartment_type = request.form.get("apartment", "").strip()
    payment_method = request.form.get("payment", "").strip()
    payment_date = request.form.get("date", "").strip()
    total_days = parse_float(request.form.get("total_days"))
    reservation_fee = parse_float(request.form.get("reservation_fee"))
    amount_paid = parse_float(request.form.get("amount_paid"))

    apartment_type = f"{apartment_type} bedroom"
    days_cost = reservation_fee * total_days
    vat = days_cost * 0.075
    total = days_cost + vat
    balance = total - amount_paid
    date = datetime.today().strftime("%Y-%m-%d")

    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    receipt_no = next_id_code(ws)
    ws.append([
        receipt_no, name, email, apartment_type, place,
        date_coming, date_going, total_days, reservation_fee,
        vat, total, amount_paid, payment_method, payment_date, balance
    ])
    auto_adjust_columns(ws)
    wb.save(FILENAME)

    record = {
        "Receipt No": receipt_no,
        "Date": date,
        "Guest Name": name,
        "Contact/Email": email,
        "Apartment Type": apartment_type,
        "Place / Location": place,
        "Check-In Date": date_coming,
        "Check-Out Date": date_going,
        "Number of Nights": total_days,
        "Rate per Night": f"{reservation_fee:.2f}",
        "VAT": f"{vat:.2f}",
        "Total Amount": f"{total:.2f}",
        "Amount Paid": f"{amount_paid:.2f}",
        "Payment Method": payment_method,
        "Payment Date": payment_date,
        "Balance": f"{balance:.2f}"
    }

    guest_clean = sanitize_filename(name)
    output_path = os.path.join(CUSTOMER_FOLDER, f"CustomerRequest_{guest_clean}.docx")
    doc = Document()
    doc.add_heading("Reservation Request", 0)
    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(output_path)

    flash(f"Reservation submitted! Receipt No: {receipt_no}")
    return redirect(url_for("customer_form"))

@app.route("/index")
def index():
    if not logged_in():
        return redirect(url_for("login"))
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        rows.append(dict(zip([c.value for c in ws[1]], row)))
    return render_template_string(index_template, SITE_NAME=SITE_NAME, rows=rows)

@app.route("/edit/<receipt_no>")
def edit(receipt_no):
    if not logged_in():
        return redirect(url_for("login"))
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    record = None
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] == receipt_no:
            record = dict(zip(headers, row))
            break
    if not record:
        return "Record not found", 404
    return render_template_string(edit_template, record=record)

@app.route("/update/<receipt_no>", methods=["POST"])
def update(receipt_no):
    if not logged_in():
        return redirect(url_for("login"))
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    headers = [c.value for c in ws[1]]
    updated = False
    for idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
        if row[0].value == receipt_no:
            for i, header in enumerate(headers):
                val = request.form.get(header, "")
                if row[i].data_type == 'n':
                    try:
                        row[i].value = float(val)
                    except:
                        row[i].value = val
                else:
                    row[i].value = val
            updated = True
            break
    if updated:
        wb.save(FILENAME)
    return redirect(url_for("index"))

@app.route("/download")
def download():
    return send_from_directory(DATA_FOLDER, os.path.basename(FILENAME), as_attachment=True)

@app.route("/search")
def search():
    if not logged_in():
        return redirect(url_for("login"))
    query = request.args.get("q", "").lower()
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    headers = [c.value for c in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if query in str(row[1]).lower():
            rows.append(dict(zip(headers, row)))
    return render_template_string(search_template, rows=rows)

@app.route("/receipt/<receipt_no>/word")
def download_word_receipt(receipt_no):
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    headers = [c.value for c in ws[1]]
    record = None
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] == receipt_no:
            record = dict(zip(headers, row))
            break
    if not record:
        return "Receipt not found", 404
    filepath = generate_word_receipt(receipt_no, record)
    return send_from_directory(DATA_FOLDER, os.path.basename(filepath), as_attachment=True)

@app.route("/invoice/<receipt_no>/download")
def download_invoice(receipt_no):
    wb = load_workbook(FILENAME)
    ws = wb["Reservations"]
    headers = [c.value for c in ws[1]]
    record = None
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] == receipt_no:
            record = dict(zip(headers, row))
            break
    if not record:
        return "Invoice not found", 404
    filepath = generate_invoice(receipt_no, record)
    return send_from_directory(DATA_FOLDER, os.path.basename(filepath), as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
